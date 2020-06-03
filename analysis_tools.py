"""
This code is for getting statistics of evaluated summaries and saving to .docx files
"""
from preprocessing_tools import create_dir
from docx import Document
from docx.shared import Pt
import os
import pandas as pd
from preprocessing_tools import preprocess_path

#dict for printing
score_names = {
    'f': 'F - ',
    'p': 'Precision - ',
    'r': 'Recall - '
}

#columns in estimated dataframe
est_columns = [
    ('rouge-1_f',	'rouge-1_p',	'rouge-1_r'),
    ('rouge-2_f',	'rouge-2_p',	'rouge-2_r'),
    ('rouge-l_f',	'rouge-l_p',	'rouge-l_r')
]


def read_data(path_to_read):
    """Reads dataframes

    Parameters
    ----------
    path_to_read : str
        Path to read dataframes from

    Return
    ----------
    res_dict: dict
    res_dfs: list of dataframes
    """

    path_to_read = preprocess_path(path_to_read)
    folders = os.listdir(path_to_read)

    res_dict = dict()
    res_dfs = []

    for folder in folders:
        cur_df = pd.DataFrame()
        types_list = []

        folder += '/'
        files = os.listdir(path_to_read + folder)

        for file in files:
            path_to_file = path_to_read + folder + file
            df = pd.read_csv(path_to_file, sep='|', encoding='utf8')
            cur_df = pd.concat([cur_df, df], ignore_index=True)
            d_key = file[:-4]
            if d_key in res_dict.keys():
                res_dict[d_key].append((df, folder))
            else:
                res_dict[d_key] = [(df, folder[:-1])]

            types_list += [d_key] * df.shape[0]

        cur_df['type'] = types_list
        res_dfs.append((cur_df, folder[:-1]))

    return res_dict, res_dfs


def get_statistics_str(df):
    """Counts statistics in dataframe

    Parameters
    ----------
    df : dataframe
        dataframe to analyze
    """

    s = ''
    for est_cols in est_columns:
        s += est_cols[0].split('_')[0].upper() + ':'
        for col in est_cols:
            s += '\t' + score_names[col.split('_')[1]]
            s += str(df[col].mean())[:6]
        s += '\n'
    key_words_est = str(df['key_word_part'].mean())
    s += '\nKEY_WORDS: ' + key_words_est[:min(6, len(key_words_est))]
    s += '\n======================================'
    return s


def create_df_info(df, name, path_to_save, possible_types):
    """Counts statistics in dataframe and save it to file

    Parameters
    ----------
    df : dataframe
        dataframe to analyze
    name: str
        name of file to save statistics in
    path_to_save: str
        path to folder for saving statistics file in
    possible_types: list
        list of types of texts under analysis
    """

    path_to_save = preprocess_path(path_to_save)
    f = open(path_to_save + name, 'w', encoding='utf8')

    f.write('Статистика по всем типам текстов:' + '\n\n')
    f.write(get_statistics_str(df) + '\n\n')

    for data_type in possible_types:
        f.write('Статистика по типу {}:'.format(data_type) + '\n\n')
        s = get_statistics_str(df[df['type'] == data_type])
        f.write(s + '\n\n')
    f.close()


def create_dfs_info(dfs, path_to_save, possible_types):
    """Counts statistics in dataframes and save it to files

    Parameters
    ----------
    dfs : list
        list of dataframes to analyze
    path_to_save: str
        path to folder for saving statistics file in
    possible_types: list
        list of types of texts under analysis
    """
    path_to_save = preprocess_path(path_to_save)
    create_dir(path_to_save)
    for df_name in dfs:
        create_df_info(df_name[0], df_name[1] + '.txt', path_to_save, possible_types)


def create_doc_file(titles, texts, save_folder, name):
    """Creates .docx file

    Parameters
    ----------
    titles : list
        list of texts' titles
    texts: list
        list of texts
    save_folder: str
        folder to save file
    name: str
        name of file being created
    """
    document = Document()
    font = document.styles['Normal'].font
    font.name = 'Times New Roman'
    font.size = Pt(14)

    document.add_heading(texts[0], 0)

    for i in range(1, len(titles)-1, 2):
        p = document.add_paragraph('')
        p.add_run(titles[i]).bold = True
        p.add_run(texts[i].replace('\n', ' '))
        p = document.add_paragraph('')
        p.add_run(titles[i+1]).bold = True
        p.add_run(texts[i+1].replace('\n', ' '))
        document.add_paragraph('=================================')

    p = document.add_paragraph('')
    p.add_run(titles[-1]).bold = True
    p.add_run(texts[-1].replace('\n', ' '))

    document.save(save_folder + name)


def create_topic_files(list_dfs, save_folder):
    """Counts statistics, creating .docx files

    Parameters
    ----------
    list_dfs : list
        list of dataframes
    save_folder: str
        folder to save file
    """
    doc_cnt = list_dfs[0][0].shape[0]
    for i in range(doc_cnt):
        titles = []
        texts = []

        titles.append('Название:')
        doc_name = list_dfs[0][0]['title'][i]
        texts.append(doc_name)

        titles.append('Человек: ')
        texts.append(list_dfs[0][0]['annotation'][i])

        titles.append('Key words: ')
        texts.append(list_dfs[0][0]['key_words'][i])

        for df_method in list_dfs:
            titles.append(df_method[1] + ': ')
            texts.append(df_method[0]['generated'][i])
            titles.append('Key words part: ')
            texts.append(str(df_method[0]['key_word_part'][i]))

        titles.append('Текст: ')
        texts.append(list_dfs[0][0]['text'][i])

        create_doc_file(titles, texts, save_folder, str(i) + '.docx')


def process_topics(res_dict, path_to_save):
    """Full processing, getting statistics, saving

    Parameters
    ----------
    res_dict : dict
        dict of possible types and dataframes
    path_to_save: str
        folder to save files
    """
    path_to_save = preprocess_path(path_to_save)
    for file_type in res_dict.keys():
        save_folder = path_to_save + file_type + '/'
        create_dir(save_folder)
        create_topic_files(res_dict[file_type], save_folder)

